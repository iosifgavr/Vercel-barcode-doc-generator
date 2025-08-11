from flask import Flask, request, send_file, render_template_string
from docx import Document
from docx.shared import Mm, Pt
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from io import BytesIO
import barcode
from barcode.writer import ImageWriter
from PIL import Image

app = Flask(__name__)

HTML = """
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <link rel="icon" href="/static/favicon.ico" type="image/x-icon">
    <title>G101 Î”Î·Î¼Î¹Î¿Ï…ÏÎ³Î¯Î± Barcode</title>
    <style>
        body {
            font-family: sans-serif;
            max-width: 800px;
            margin: auto;
            padding: 20px;
            background-image: url('/static/background.jpg');
            background-size: cover;
            background-repeat: no-repeat;
            background-attachment: fixed;
        }
        #logo {
            position: fixed;
            top: 10px;
            left: 10px;
            width: 200px;
            height: auto;
            z-index: 1000;
        }
        button {
            background-color: #007BFF; 
            color: white;              
            border: none;
            padding: 10px 15px;
            cursor: pointer;
            border-radius: 4px;
            font-weight: bold;
            transition: background-color 0.3s ease;
        }
        button:hover {
            background-color: #0056b3;
        }
        input { margin: 5px 0; width: 100%; padding: 8px; }
        table { width: 100%; border-collapse: collapse; margin-top: 20px; background-color: #f5f5f5; border: 1px solid #ccc;  }
        th, td { border: 1px solid #ccc; padding: 8px; text-align: left; background-color: white;}
        button { padding: 10px 15px; margin-top: 10px; }
        td > button { margin-right: 5px; }
        h2 { text-align: center; }
        
        #popup {
            position: fixed;
            top: 50%; left: 50%;
            transform: translate(-50%, -50%);
            background: white;
            border: 2px solid #007BFF;
            padding: 20px 30px;
            box-shadow: 0 0 10px rgba(0,0,0,0.25);
            z-index: 2000;
            text-align: center;
            font-weight: bold;
            font-size: 18px;
        }
    </style>
</head>
<body>

<div id="popup">
    <h1 style="color: red;">Î Î¡ÎŸÎ£ÎŸÎ§Î—</h1>
    <p>Î— Ï€Î±ÏÎ¿ÏÏƒÎ± Î¹ÏƒÏ„Î¿ÏƒÎµÎ»Î¯Î´Î± Î´ÎµÎ½ Î±Ï€Î¿Ï„ÎµÎ»ÎµÎ¯ ÎµÏ€Î¯ÏƒÎ·Î¼Î· Î® ÎµÎ³ÎºÎµÎºÏÎ¹Î¼Î­Î½Î· Ï€Î»Î±Ï„Ï†ÏŒÏÎ¼Î± Ï„Î·Ï‚ ÎµÏ„Î±Î¹ÏÎµÎ¯Î±Ï‚ Î’. ÎšÎ±Ï…ÎºÎ¬Ï‚ Î‘.Î•..
Î”ÎµÎ½ Ï†Î­ÏÎµÎ¹ ÎºÎ±Î¼Î¯Î± Î½Î¿Î¼Î¹ÎºÎ® ÎµÏ…Î¸ÏÎ½Î· Î® Î´Î¹ÎºÎ±Î¹ÏÎ¼Î±Ï„Î± Ï€Î½ÎµÏ…Î¼Î±Ï„Î¹ÎºÎ®Ï‚ Î¹Î´Î¹Î¿ÎºÏ„Î·ÏƒÎ¯Î±Ï‚ Î±Ï€ÏŒ Ï„Î·Î½ ÎµÏ„Î±Î¹ÏÎµÎ¯Î±.
Î¤Î¿ Ï€ÎµÏÎ¹ÎµÏ‡ÏŒÎ¼ÎµÎ½Î¿ ÎºÎ±Î¹ Î¿Î¹ Î»ÎµÎ¹Ï„Î¿Ï…ÏÎ³Î¯ÎµÏ‚ Ï„Î·Ï‚ Î¹ÏƒÏ„Î¿ÏƒÎµÎ»Î¯Î´Î±Ï‚ Ï€Î±ÏÎ­Ï‡Î¿Î½Ï„Î±Î¹ Î±Ï€Î¿ÎºÎ»ÎµÎ¹ÏƒÏ„Î¹ÎºÎ¬ Î³Î¹Î± ÎµÎ½Î·Î¼ÎµÏÏ‰Ï„Î¹ÎºÎ¿ÏÏ‚ ÎºÎ±Î¹ Î´Î¿ÎºÎ¹Î¼Î±ÏƒÏ„Î¹ÎºÎ¿ÏÏ‚ ÏƒÎºÎ¿Ï€Î¿ÏÏ‚.
Î— Ï‡ÏÎ®ÏƒÎ· Ï„Î·Ï‚ Î³Î¯Î½ÎµÏ„Î±Î¹ Ï…Ï€ÏŒ Ï„Î·Î½ Î±Ï€Î¿ÎºÎ»ÎµÎ¹ÏƒÏ„Î¹ÎºÎ® ÎµÏ…Î¸ÏÎ½Î· Ï„Î¿Ï… Ï‡ÏÎ®ÏƒÏ„Î·.</p>
    <button id="closePopup">ÎšÎ±Ï„Î±Î½Î¿Ï ÎºÎ±Î¹ Î±Ï€Î¿Î´Î­Ï‡Î¿Î¼Î±Î¹</button>
</div>

<img src="/static/logo.png" alt="Logo" id="logo" />
<h2>ÎšÎ±Ï„Î±Ï‡ÏÏÎ¹ÏƒÎ· Î ÏÎ¿ÏŠÏŒÎ½Ï„Ï‰Î½</h2>
<form id="productForm">
    <input type="text" id="barcode" placeholder="Barcode" required><br>
    <input type="text" id="description" placeholder="Î ÎµÏÎ¹Î³ÏÎ±Ï†Î®" required><br>
    <input type="text" id="code" placeholder="7ÏˆÎ®Ï†Î¹Î¿Ï‚ ÎšÏ‰Î´Î¹ÎºÏŒÏ‚ SAP" maxlength="7" required><br>
    <input type="text" id="manufacturer_part" placeholder="Î‘Ï. Î•Î¾Î±ÏÏ„Î®Î¼Î±Ï„Î¿Ï‚ ÎšÎ±Ï„Î±ÏƒÎºÎµÏ…Î±ÏƒÏ„Î®" required><br>
    <input type="text" id="supplier" placeholder="Î ÏÎ¿Î¼Î·Î¸ÎµÏ…Ï„Î®Ï‚" required><br>
    <button type="submit">Î ÏÎ¿ÏƒÎ¸Î®ÎºÎ·</button>
    <button type="button" id="clearAllBtn" style="margin-left:10px; background-color: red;">Î”Î¹Î±Î³ÏÎ±Ï†Î® ÎŒÎ»Ï‰Î½</button>
</form>
<table id="productsTable">
    <thead>
        <tr>
            <th>Barcode</th>
            <th>Î ÎµÏÎ¹Î³ÏÎ±Ï†Î®</th>
            <th>ÎšÏ‰Î´Î¹ÎºÏŒÏ‚ SAP</th>
            <th>Î‘Ï. Î•Î¾Î±ÏÏ„Î®Î¼Î±Ï„Î¿Ï‚ ÎšÎ±Ï„Î±ÏƒÎºÎµÏ…Î±ÏƒÏ„Î®</th>
            <th>Î ÏÎ¿Î¼Î·Î¸ÎµÏ…Ï„Î®Ï‚</th>
            <th>Î•Î½Î­ÏÎ³ÎµÎ¹ÎµÏ‚</th>
        </tr>
    </thead>
    <tbody></tbody>
</table>
<button onclick="downloadDoc()">ÎšÎ±Ï„ÎµÎ²Î¬ÏƒÏ„Îµ .doc</button>
<script>

const popup = document.getElementById('popup');
const closeBtn = document.getElementById('closePopup');
closeBtn.onclick = () => {
    popup.style.display = 'none';
};

const form = document.getElementById('productForm');
const table = document.getElementById('productsTable').querySelector('tbody');
const products = [];
let editIndex = -1;

form.onsubmit = function(e) {
    e.preventDefault();
    const barcode = document.getElementById('barcode').value;
    const description = document.getElementById('description').value;
    const code = document.getElementById('code').value;
    const manufacturer_part = document.getElementById('manufacturer_part').value;
    const supplier = document.getElementById('supplier').value;

    if (editIndex === -1) {
        products.push({ barcode, description, code, manufacturer_part, supplier });
    } else {
        products[editIndex] = { barcode, description, code, manufacturer_part, supplier };
        editIndex = -1;
    }

    updateTable();
    form.reset();
};

function updateTable() {
    table.innerHTML = '';
    products.forEach((item, index) => {
        const row = document.createElement('tr');
        row.innerHTML = `
            <td>${item.barcode}</td>
            <td>${item.description}</td>
            <td>${item.code}</td>
            <td>${item.manufacturer_part}</td>
            <td>${item.supplier}</td>
            <td>
                <button onclick="editProduct(${index})">âœï¸</button>
                <button onclick="deleteProduct(${index})">ğŸ—‘ï¸</button>
            </td>`;
        table.appendChild(row);
    });
}

function editProduct(index) {
    const product = products[index];
    document.getElementById('barcode').value = product.barcode;
    document.getElementById('description').value = product.description;
    document.getElementById('code').value = product.code;
    document.getElementById('manufacturer_part').value = product.manufacturer_part;
    document.getElementById('supplier').value = product.supplier;
    editIndex = index;
}

function deleteProduct(index) {
    products.splice(index, 1);
    updateTable();
}

function downloadDoc() {
    fetch('/generate_doc', {
        method: 'POST',
        headers: { 'Content-Type': 'application/json' },
        body: JSON.stringify({ products })
    })
    .then(response => {
        if (!response.ok) throw new Error("Server error");
        return response.blob();
    })
    .then(blob => {
        const url = window.URL.createObjectURL(blob);
        const a = document.createElement('a');
        a.href = url;
        a.download = 'products.docx';
        a.click();
        window.URL.revokeObjectURL(url);
    })
    .catch(e => alert(e.message));
}

document.getElementById('clearAllBtn').onclick = function() {
    if (confirm("Î˜Î­Î»ÎµÎ¹Ï‚ ÏƒÎ¯Î³Î¿Ï…ÏÎ± Î½Î± Î´Î¹Î±Î³ÏÎ¬ÏˆÎµÎ¹Ï‚ ÏŒÎ»Î± Ï„Î± Ï€ÏÎ¿ÏŠÏŒÎ½Ï„Î±;")) {
        products.length = 0; // Î±Î´ÎµÎ¹Î¬Î¶ÎµÎ¹ Ï„Î¿Î½ Ï€Î¯Î½Î±ÎºÎ±
        updateTable();
    }
};

</script>
</body>
</html>
"""

def set_font(run, name, size_pt):
    run.font.size = Pt(size_pt)
    run.font.name = name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)

@app.route('/')
def index():
    return render_template_string(HTML)

@app.route('/generate_doc', methods=['POST'])
def generate_doc():
    data = request.json
    products = data.get('products', [])

    doc = Document()
    section = doc.sections[-1]
    section.page_height = Mm(150)
    section.page_width = Mm(100)
    section.orientation = WD_ORIENT.PORTRAIT
    section.top_margin = Mm(3)
    section.left_margin = Mm(10)
    section.right_margin = Mm(10)
    section.bottom_margin = Mm(1)

    for idx, item in enumerate(products):
        if idx > 0:
            doc.add_page_break()

        # Î”Î·Î¼Î¹Î¿Ï…ÏÎ³Î¯Î± barcode
        barcode_stream = BytesIO()
        code128 = barcode.get('code128', item['barcode'], writer=ImageWriter())
        code128.write(barcode_stream)
        barcode_stream.seek(0)

        img = Image.open(barcode_stream).copy()
        img_buffer = BytesIO()
        img.save(img_buffer, format="PNG")
        img_buffer.seek(0)

        # 1) Barcode 8x7 cm
        barcode_paragraph = doc.add_paragraph()
        barcode_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        barcode_run = barcode_paragraph.add_run()
        barcode_run.add_picture(img_buffer, width=Mm(80), height=Mm(70))

        # 2) Î ÎµÏÎ¹Î³ÏÎ±Ï†Î® Î¼Îµ Calibri ÎºÎ±Î¹ 20pt
        desc_paragraph = doc.add_paragraph(item['description'])
        desc_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in desc_paragraph.runs:
            set_font(run, "Calibri", 18)

        # 3) ÎšÏ‰Î´Î¹ÎºÏŒÏ‚ SAP
        code_paragraph = doc.add_paragraph(f"ÎšÎ©Î”Î™ÎšÎŸÎ£ SAP: {item['code']}")
        code_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in code_paragraph.runs:
            set_font(run, "Calibri", 20)

        # 4) Î‘Ï. Î•Î¾Î±ÏÏ„Î®Î¼Î±Ï„Î¿Ï‚ ÎšÎ±Ï„Î±ÏƒÎºÎµÏ…Î±ÏƒÏ„Î®
        manufacturer_paragraph = doc.add_paragraph(f"MPN: {item['manufacturer_part']} {item['supplier']}")
        manufacturer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in manufacturer_paragraph.runs:
            set_font(run, "Calibri", 14)

        # 5) Î ÏÎ¿Î¼Î·Î¸ÎµÏ…Ï„Î®Ï‚
#        supplier_paragraph = doc.add_paragraph(f"Î Î¡ÎŸÎœÎ—Î˜Î•Î¥Î¤Î—Î£: {item['supplier']}")
 #       supplier_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
  #      for run in supplier_paragraph.runs:
   #         set_font(run, "Calibri", 20

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return send_file(
        buffer,
        as_attachment=True,
        download_name='products.docx',
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )

if __name__ == '__main__':
    app.run(debug=False)
